# coding: utf-8
"""东莞致嘉 - 自动文档生成（Streamlit 版）"""
import os
import re
import shutil
import tempfile
import traceback
from copy import copy  # noqa: F401
from datetime import datetime

import pandas as pd
import streamlit as st
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt
from openpyxl import load_workbook, Workbook  # noqa: F401
from openpyxl.styles import Border, Font, PatternFill  # noqa: F401
from openpyxl.styles import Alignment, Side
from openpyxl.utils import get_column_letter  # noqa: F401

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
STATIC_DIR = os.path.join(BASE_DIR, "static", "zhijia")

# --- 全局模板定义 ---
ITEM_DECLARATION_TEMPLATES = {
    "五金冲压模具": {
        "code": "8207300090",
        "template_lines": [
            "五金冲压模具 8207300090",
            "1.     品牌类型:无品牌 2.出口享惠情况:不享惠  3.用途:冲压用  4.材质:钢铁制  5.种类:冲压模 6.品牌:无牌  7.是否带有工作部件:否  8.型号: {model}",
        ],
    },
    "机械手自动传送机/用于物料传送": {
        "code": "8428909090",
        "template_lines": [
            "机械手自动传送机/用于物料传  8428909090",
            "1.     品牌类型:无品牌 2.出口享惠情况:不享惠  3.品牌:无牌 4.用途:用于物料传送 5.型号: {model}",
        ],
    },
    "检具": {
        "code": "9031809090",
        "template_lines": [
            "检具 9031809090",
            "1.品牌类型:无品牌 2.出口享惠情况:不享惠 3.用途:检测模具加工出产品的精密度 4.原理: 通过CNC加工产品的3D模型与该批模具生产的产品进行检测，测量   5.功能: 用于测量,检验模具样品用  6.品牌:无牌  7.型号: {model}",
        ],
    },
    "总成检具": {
        "code": "9031809090",
        "template_lines": [
            "总成检具 9031809090",
            "1.品牌类型:无品牌 2.出口享惠情况:不享惠 3.用途:检测模具加工出产品的精密度 4.原理: 通过CNC加工产品的3D模型与该批模具生产的产品进行检测，测量   5.功能: 用于测量,检验模具样品用  6.品牌:无牌  7.型号: {model}",
        ],
    },
    "检具推车": {
        "code": "8716800000",
        "template_lines": [
            "检具推车  8716800000",
            "1.品牌类型:无品牌    2.出口享惠情况:不享惠 .  3.型号: {model}",
        ],
    },
    "五金冲压模具配件/冲头.入子": {
        "code": "8207300090",
        "template_lines": [
            "8207300090五金冲压模具配件/冲头.入子",
            "1.品牌类型:无品牌 2.出口享惠情况:不享惠  3.用途: 五金冲压模具用  4.材质:钢铁制  5.种类: 冲头.入子 6.品牌:无牌  7.是否带有工作部件:否  8.型号: {model}",
        ],
    },
    "汽车五金配件/用于支架系统": {
        "code": "8708299000",
        "template_lines": [
            "汽车五金配件/用于支架系统  8708299000",
            "1.品牌类型:无品牌 2.出口享惠情况:不享惠  3.品牌:无牌 4.适用车型:通用 5.零部件编号: {model}",
        ],
    },
}


# --- 辅助函数 ---
def get_file_naming_date_str(date_input):
    date_input = str(date_input).strip()
    formats_to_try = ["%Y-%m-%d", "%Y/%m/%d", "%Y.%m.%d", "%Y%m%d", "%m-%d-%Y", "%m/%d/%Y", "%m.%d.%Y", "%Y年%m月%d日"]
    for fmt in formats_to_try:
        try:
            return datetime.strptime(date_input, fmt).strftime("%Y%m%d")
        except ValueError:
            continue
    return datetime.now().strftime("%Y%m%d")


def format_date(date_str):
    try:
        date_str = str(date_str).strip()
        if "年" in date_str and "月" in date_str:
            return date_str
        formats = ["%Y-%m-%d", "%Y/%m/%d", "%Y.%m.%d", "%Y%m%d", "%m-%d-%Y", "%m/%d/%Y", "%m.%d.%Y"]
        for fmt in formats:
            try:
                return datetime.strptime(date_str, fmt).strftime("%Y年%m月%d日")
            except ValueError:
                continue
        return date_str
    except Exception:
        return str(date_str)


def get_item_category(chinese_name):
    """根据发票中提取的中文品名，匹配对应的申报模版分类"""
    if "\n" in chinese_name:
        parts = [p.strip() for p in chinese_name.split("\n") if p.strip()]
        main_name = parts[1] if len(parts) > 1 and re.search(r'[\u4e00-\u9fa5]', parts[1]) else parts[0]
    else:
        main_name = chinese_name
        
    lower_name = chinese_name.lower()
    main_lower = main_name.lower()

    if "机械手" in main_lower or "传送机" in main_lower or "robot" in lower_name:
        return "机械手自动传送机/用于物料传送"
    elif "总成检具" in main_lower or "assembly fixture" in lower_name:
        return "总成检具"
    elif "冲头" in main_lower or "入子" in main_lower or ("模具配件" in main_lower):
        return "五金冲压模具配件/冲头.入子"
    elif "推车" in main_lower or "cart" in lower_name or ("检具" in main_lower and "车" in main_lower):
        return "检具推车"
    elif "检具" in main_lower:
        return "检具"
    elif "冲压模具" in main_lower or "stamping die" in lower_name:
        return "五金冲压模具"
    elif "汽车五金配件" in main_lower or "支架系统" in main_lower:
        return "汽车五金配件/用于支架系统"
        
    return None


def get_unit_str(chinese_name):
    """根据品名自动识别并补充计量单位"""
    if "\n" in chinese_name:
        parts = [p.strip() for p in chinese_name.split("\n") if p.strip()]
        main_name = parts[1] if len(parts) > 1 and re.search(r'[\u4e00-\u9fa5]', parts[1]) else parts[0]
    else:
        main_name = chinese_name
    lower_name = main_name.lower()
    
    if "推车" in lower_name or "cart" in lower_name:
        return "辆"
    elif "机械手" in lower_name or "传送机" in lower_name or "robot" in lower_name:
        return "台"
    elif "配件" in lower_name or "支架" in lower_name or "冲头" in lower_name or "入子" in lower_name:
        return "个"
    elif "检具" in lower_name or "模具" in lower_name or "die" in lower_name:
        return "套"
    return "套"


def read_invoice_data(file_path):
    """
    从Excel文件读取发票数据，包括商品信息、毛重净重、总金额及箱数。
    """
    try:
        if not os.path.exists(file_path):
            return None

        file_ext = os.path.splitext(file_path)[1].lower()

        target_sheet = None
        xls = pd.ExcelFile(file_path)

        if "发票" in xls.sheet_names:
            target_sheet = "发票"
        else:
            for sheet in xls.sheet_names:
                if "发票" in sheet:
                    target_sheet = sheet
                    break

        if not target_sheet:
            for sheet in xls.sheet_names:
                if sheet.lower() in ["invoice", "装箱单", "sheet1", "sheet2"]:
                    target_sheet = sheet
                    break
        target_sheet = target_sheet or xls.sheet_names[0]

        if file_ext == ".xls":
            invoice_df = pd.read_excel(file_path, sheet_name=target_sheet)
            wb = None
            ws = None
            merged_ranges = []
        else:
            try:
                wb = load_workbook(file_path, data_only=True)
                ws = wb[target_sheet]
                merged_ranges = list(ws.merged_cells.ranges)
                data = []
                for row in ws.iter_rows(values_only=True):
                    data.append([cell if cell is not None else "" for cell in row])
                invoice_df = pd.DataFrame(data)
            except Exception:
                invoice_df = pd.read_excel(file_path, sheet_name=target_sheet)
                wb = None
                ws = None
                merged_ranges = []

        header_row_idx = 0
        header_keywords = ["中文品名", "品名", "Description", "产品名称"]
        for i in range(min(20, len(invoice_df))):
            row_str = " ".join([str(x) for x in invoice_df.iloc[i].values if not pd.isna(x) and str(x) != ""])
            if any(k in row_str for k in header_keywords):
                header_row_idx = i
                break

        data_start_row = header_row_idx + 1

        company_name, date_str, destination, contract_number = None, None, None, None

        for i in range(min(header_row_idx, 10)):
            for j in range(min(10, len(invoice_df.columns))):
                val = str(invoice_df.iloc[i, j]).strip()
                if not val or val.lower() == "nan" or val == "":
                    continue
                if company_name is None and ("公司" in val or "LTD" in val) and len(val) > 4:
                    company_name = val.replace("To:", "").strip()
                if date_str is None and ("日期" in val or "Date" in val):
                    parts = re.split(r"[:：]", val)
                    candidate = parts[-1].strip() if len(parts) > 1 else val
                    if len(candidate) < 5 and j + 1 < len(invoice_df.columns):
                        candidate = str(invoice_df.iloc[i, j + 1]).strip()
                    if re.search(r"\d", candidate):
                        date_str = candidate
                if destination is None and ("目的国" in val or "Destination" in val):
                    parts = re.split(r"[:：]", val)
                    candidate = parts[-1].strip() if len(parts) > 1 else val
                    if len(candidate) < 2 and j + 1 < len(invoice_df.columns):
                        candidate = str(invoice_df.iloc[i, j + 1]).strip()
                    destination = candidate

        if contract_number is None or str(contract_number).lower() == "nan" or contract_number == "":
            header_vals = [str(x).strip() for x in invoice_df.iloc[header_row_idx]]
            contract_col = -1
            for idx, h in enumerate(header_vals):
                if any(k in h for k in ["合同号", "Contract No", "NO."]):
                    contract_col = idx
                    break
            if contract_col != -1 and header_row_idx + 1 < len(invoice_df):
                contract_number = str(invoice_df.iloc[header_row_idx + 1, contract_col]).strip()
            else:
                for i in range(min(header_row_idx, 10)):
                    for j in range(min(10, len(invoice_df.columns))):
                        val = str(invoice_df.iloc[i, j]).strip()
                        if "合同号" in val or "Contract No" in val:
                            parts = re.split(r"[:：]", val)
                            if len(parts) > 1:
                                contract_number = parts[-1].strip()
                            elif j + 1 < len(invoice_df.columns):
                                contract_number = str(invoice_df.iloc[i, j + 1]).strip()
                            break
                    if contract_number:
                        break

        col_map = {}
        header_vals = [str(x).strip() for x in invoice_df.iloc[header_row_idx]]

        kw_map = {
            "desc": ["中文品名", "品名", "Description", "货名"],
            "model": ["型号", "Model", "Part No"],
            "qty": ["数量", "Quantity", "QTY"],
            "price": ["单价", "Unit Price"],
            "amt": ["金额", "Amount", "Total"],
        }
        for k, v in kw_map.items():
            for kw in v:
                if kw in header_vals:
                    col_map[k] = header_vals.index(kw)
                    break

        desc_col = col_map.get("desc", 0)
        model_col = col_map.get("model", None)
        qty_col = col_map.get("qty", 4)
        price_col = col_map.get("price", 7)
        amt_col = col_map.get("amt", 8)

        if model_col is None:
            for j in range(len(header_vals)):
                if j in [desc_col, qty_col, price_col, amt_col]:
                    continue
                count = 0
                for r in range(data_start_row, min(data_start_row + 5, len(invoice_df))):
                    val = str(invoice_df.iloc[r, j]).strip()
                    if len(val) > 3 and re.search(r"[A-Za-z]\d", val):
                        count += 1
                if count > 0:
                    model_col = j
                    break

        temp_total_amount = 0.0
        net_weight = 0.0
        gross_weight = 0.0
        found_weights = False

        temp_data_storage = []

        total_keywords = ["总计", "合计", "total", "TOTAL", "合计金额"]
        net_weight_keywords = ["净重", "Net Weight", "NET", "N.W.", "N.W"]
        gross_weight_keywords = ["毛重", "Gross Weight", "G.W.", "G.W", "Gross"]

        total_row = -1
        net_weight_col = -1
        gross_weight_col = -1

        for i in range(len(invoice_df)):
            row = invoice_df.iloc[i]
            for col_idx, cell_val in enumerate(row):
                cell_str = str(cell_val).strip()
                if any(keyword in cell_str for keyword in total_keywords):
                    total_row = i
                    break
            if total_row != -1:
                break

        if total_row == -1:
            total_row = len(invoice_df) - 1

        for i in range(min(20, len(invoice_df))):
            row = invoice_df.iloc[i]
            for col_idx, cell_val in enumerate(row):
                cell_str = str(cell_val).strip()
                if any(keyword in cell_str for keyword in net_weight_keywords):
                    net_weight_col = col_idx
                    break
            if net_weight_col != -1:
                break

        for i in range(min(20, len(invoice_df))):
            row = invoice_df.iloc[i]
            for col_idx, cell_val in enumerate(row):
                cell_str = str(cell_val).strip()
                if any(keyword in cell_str for keyword in gross_weight_keywords):
                    gross_weight_col = col_idx
                    break
            if gross_weight_col != -1:
                break

        if total_row != -1 and total_row < len(invoice_df):
            total_row_data = invoice_df.iloc[total_row]

            if net_weight_col != -1 and net_weight_col < len(total_row_data):
                net_val = str(total_row_data.iloc[net_weight_col]).strip()
                if net_val and net_val.lower() != "nan":
                    numbers = re.findall(r"[-+]?\d*\.?\d+", net_val)
                    if numbers:
                        net_weight = float(numbers[0])

            if gross_weight_col != -1 and gross_weight_col < len(total_row_data):
                gross_val = str(total_row_data.iloc[gross_weight_col]).strip()
                if gross_val and gross_val.lower() != "nan":
                    numbers = re.findall(r"[-+]?\d*\.?\d+", gross_val)
                    if numbers:
                        gross_weight = float(numbers[0])

            if net_weight == 0.0 or gross_weight == 0.0:
                numbers_found = []
                for col_idx in range(len(total_row_data)):
                    val = str(total_row_data.iloc[col_idx]).strip()
                    if val and val.lower() != "nan":
                        nums = re.findall(r"[-+]?\d*\.?\d+", val)
                        for num in nums:
                            try:
                                num_val = float(num)
                                if 0 < num_val < 10000:
                                    numbers_found.append((col_idx, num_val, val))
                            except Exception:
                                pass

                if len(numbers_found) >= 2:
                    if net_weight == 0.0:
                        net_weight = numbers_found[0][1]
                    if gross_weight == 0.0:
                        gross_weight = numbers_found[1][1]
                elif len(numbers_found) == 1 and gross_weight == 0.0:
                    gross_weight = numbers_found[0][1]

        found_weights = net_weight > 0.0 or gross_weight > 0.0

        for i in range(data_start_row, len(invoice_df)):
            row = invoice_df.iloc[i]

            if i == total_row:
                continue

            first_col_val = str(row.iloc[desc_col]).strip() if desc_col < len(row) else ""
            item = [None] * 9
            item[0] = first_col_val

            if not item[0] and all(pd.isna(x) or str(x).strip() == "" for x in row.iloc[:min(len(row), 5)]):
                continue

            if model_col is not None and model_col < len(row):
                item[2] = str(row.iloc[model_col]).strip()
            if qty_col < len(row):
                item[4] = row.iloc[qty_col]
            if price_col < len(row):
                item[7] = row.iloc[price_col]
            if amt_col < len(row):
                item[8] = row.iloc[amt_col]

            try:
                if item[8] and str(item[8]).strip():
                    clean_amt = float(str(item[8]).replace(",", "").replace("EUR", "").replace("$", "").replace("€", "").replace("¥", "").strip())
                    item[8] = clean_amt
                    temp_total_amount += clean_amt
                if item[7] and str(item[7]).strip():
                    item[7] = float(str(item[7]).replace(",", "").replace("EUR", "").replace("$", "").replace("€", "").replace("¥", "").strip())
            except Exception:
                pass

            if item[0] and item[0] != "" and item[0].lower() != "nan":
                qty_valid = False
                try:
                    if isinstance(item[4], (int, float)) and item[4] > 0:
                        qty_valid = True
                    elif isinstance(item[4], str) and item[4].strip():
                        qty_float = float(item[4].replace(",", "").strip())
                        if qty_float > 0:
                            item[4] = qty_float
                            qty_valid = True
                except Exception:
                    pass

                if qty_valid:
                    temp_data_storage.append(item)

        data = temp_data_storage

        total_amount_col = -1
        total_amount_row = -1
        for i in range(len(invoice_df)):
            for j, val in enumerate(invoice_df.iloc[i]):
                val_str = str(val).strip()
                if "总价" in val_str and total_amount_col == -1:
                    total_amount_col = j
                if "总计" in val_str:
                    total_amount_row = i

        total_amount = 0.0
        if total_amount_row != -1 and total_amount_col != -1:
            try:
                intersect_val = str(invoice_df.iloc[total_amount_row, total_amount_col])
                clean_val = intersect_val.replace(",", "").replace("EUR", "").replace("$", "").replace("€", "").replace("¥", "").strip()
                nums = re.findall(r"[-+]?\d*\.?\d+", clean_val)
                if nums:
                    total_amount = float(nums[0])
            except Exception:
                pass

        if total_amount <= 0:
            total_amount = temp_total_amount

        total_packages = 0
        try:
            pack_sheet_name = None
            for sheet in xls.sheet_names:
                if "装箱单" in sheet:
                    pack_sheet_name = sheet
                    break
            
            if pack_sheet_name:
                pack_df = pd.read_excel(file_path, sheet_name=pack_sheet_name)
                box_col = -1
                total_pkg_row = -1
                
                for i in range(len(pack_df)):
                    for j, val in enumerate(pack_df.iloc[i]):
                        val_str = str(val).strip()
                        if "箱数" in val_str and box_col == -1:
                            box_col = j
                        if "总计" in val_str:
                            total_pkg_row = i
                
                if box_col != -1 and total_pkg_row != -1:
                    intersect_val = str(pack_df.iloc[total_pkg_row, box_col])
                    nums = re.findall(r'\d+', intersect_val.replace(",", ""))
                    if nums:
                        total_packages = int(nums[0])
        except Exception:
            pass

        if total_packages <= 0:
            total_packages = len(data)

        if not company_name:
            company_name = "东莞致嘉金属科技有限公司"
        if not date_str:
            date_str = datetime.now().strftime("%Y-%m-%d")
        if not destination:
            destination = "法国"
        if not contract_number or str(contract_number).lower() == "nan" or contract_number == "":
            contract_number = "ZTD" + datetime.now().strftime("%Y%m%d") + "001"
        for item in data:
            if len(item) > 1:
                item[1] = contract_number

        if wb:
            try:
                wb.close()
            except Exception:
                pass

        return {
            "company_name": company_name,
            "date": date_str,
            "destination": destination,
            "contract_number": contract_number,
            "data": data,
            "total_amount": total_amount,
            "total_packages": total_packages,
            "net_weight": net_weight,
            "gross_weight": gross_weight,
            "found_weights": found_weights,
        }
    except Exception:
        traceback.print_exc()
        return None


# --- 生成申报要素 (Word) ---
def create_declaration_elements(invoice_data, output_dir):
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(10.5)
    style.element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "Microsoft YaHei")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("申报要素")
    run.font.bold = True
    run.font.size = Pt(22)
    doc.add_paragraph()

    cn_nums = ["一", "二", "三", "四", "五", "六", "七", "八", "九", "十", "十一", "十二", "十三", "十四", "十五"]
    item_counter = 0

    for item in invoice_data["data"]:
        c_name = str(item[0]) if item[0] else ""
        model = str(item[2]) if item[2] and str(item[2]).lower() != "nan" else "无"
        cat = get_item_category(c_name)

        if cat and cat in ITEM_DECLARATION_TEMPLATES:
            item_counter += 1
            tpl = ITEM_DECLARATION_TEMPLATES[cat]
            prefix = cn_nums[item_counter - 1] if item_counter - 1 < len(cn_nums) else str(item_counter)

            p1 = doc.add_paragraph()
            p1.paragraph_format.space_after = Pt(0)
            r1 = p1.add_run(prefix + "、")
            r1.font.bold = True
            r1.font.size = Pt(12)
            r2 = p1.add_run(tpl["template_lines"][0])
            r2.font.bold = True
            r2.font.size = Pt(12)

            p2 = doc.add_paragraph()
            p2.paragraph_format.space_after = Pt(0)
            p2.add_run(tpl["template_lines"][1].format(model=model))
            doc.add_paragraph().paragraph_format.space_after = Pt(10)

    fname = f"致嘉_申报要素_{get_file_naming_date_str(invoice_data['date'])}.docx"
    fpath = os.path.join(output_dir, fname)
    doc.save(fpath)
    return fpath


# --- 修改成交确认书 (Word) ---
def modify_sales_confirmation(invoice_data, template_path, output_dir, user_inputs):
    if not os.path.exists(template_path):
        return None
    try:
        doc = Document(template_path)
        data = invoice_data["data"]
        incoterms = user_inputs.get("incoterms", "CIF")
        currency_type = user_inputs.get("currency", "EUR")
        
        buyer_name = user_inputs.get("buyer_name", "香港致达五金制品有限公司")

        currency_code_display = "EUR" if currency_type == "EUR" else "USD"
        currency_replacer = "USD" if currency_type == "USD" else "EUR"

        total_str = f"{currency_code_display} {invoice_data['total_amount']:,.2f}"
        fmt_date = format_date(invoice_data["date"])
        num_items = len(data)

        # ============== 彻底修复：使用“文本追踪法”精确定位行，杜绝错位 ==============
        target_first = "QualityNo.1"
        target_last = f"QualityNo.{num_items}"

        for t in doc.tables:
            for row in t.rows:
                # 获取整行的文本，用来判断当前是哪一件商品所在的行
                row_str = "".join(c.text for c in row.cells)
                
                # 1. 专门处理第一件商品行：如果总商品数>1，则清空最后一列原本静态的“成交方式:CIF”
                if target_first in row_str and num_items > 1:
                    last_cell = row.cells[-1] # -1 代表表格的最后一列（装运期列）
                    if "CIF" in last_cell.text or "成交方式" in last_cell.text:
                        last_cell.text = "" # 直接把这一格清空
                
                # 2. 专门处理最后一件商品行：把用户选的成交方式写入它的最后一列
                if target_last in row_str:
                    last_cell = row.cells[-1]
                    last_cell.text = f"成交方式:{incoterms}"
                    # 重新应用一下字体格式防止突兀
                    if last_cell.paragraphs:
                        for r in last_cell.paragraphs[0].runs:
                            r.font.name = "Microsoft YaHei"
                            r.font.size = Pt(10.5)
                            r.element.get_or_add_rPr().get_or_add_rFonts().set(qn('w:eastAsia'), "Microsoft YaHei")

        # ============== 构建常规替换字典 ==============
        replacements = {
            "Company1": invoice_data["company_name"],
            "香港致达五金制品有限公司": buyer_name,
            "2025年11月29日": fmt_date,
            "ZTD20251129002": invoice_data["contract_number"],
            "destinnation1": invoice_data["destination"],
            "TotalAmount1": total_str,
            "CIF": incoterms,
            "EUR": currency_replacer,
        }

        # 针对不是默认买方名称，删除预设的地址和电话
        if buyer_name != "香港致达五金制品有限公司":
            replacements["Telephone：00852-3165147"] = ""
            replacements["Telephone:00852-3165147"] = ""
            replacements["香港九龙旺角烟厂街9号兴发商业大厦2201室"] = ""

        if currency_type == "USD":
            replacements["Unit Price/EUR"] = "Unit Price/USD"
            replacements["Amount (EUR)"] = "Amount (USD)"
        else:
            replacements["Unit Price/USD"] = "Unit Price/EUR"
            replacements["Amount (USD)"] = "Amount (EUR)"

        for i in range(1, num_items + 30):
            idx = i - 1
            k_no, k_nm, k_qt, k_pr, k_am = f"QualityNo.{i}", f"Name{i}", f"Quantity{i}", f"UnitPrice{i}", f"Amount{i}"

            if i <= num_items:
                item = data[idx]
                c_name = str(item[0]) if item[0] else ""
                disp_name = c_name.split("\n")[1].strip() if "\n" in c_name and len(c_name.split("\n")) > 1 else c_name

                qty = item[4] if item[4] else 0
                if isinstance(qty, float) and qty.is_integer():
                    qty = int(qty)
                unit = get_unit_str(c_name)

                pr = item[7] if item[7] else 0.0
                am = item[8] if item[8] else 0.0
                pr_str = f"{float(pr):.2f}"
                am_str = f"{float(am):.2f}"

                replacements[k_no] = str(i)
                replacements[k_nm] = disp_name
                replacements[k_qt] = f"{qty}{unit}"
                replacements[k_pr] = pr_str
                replacements[k_am] = am_str
            else:
                for key in [k_no, k_nm, k_qt, k_pr, k_am]:
                    replacements[key] = ""

        sorted_keys = sorted(replacements.keys(), key=len, reverse=True)

        def replace_run_text_preserve_format(paragraph):
            if not paragraph.text.strip():
                return
            
            # 第一步：尝试在 run 级别安全替换（100% 保持原有格式）
            for k in sorted_keys:
                for run in paragraph.runs:
                    if k in run.text:
                        run.text = run.text.replace(k, str(replacements[k]))
            
            # 第二步：应对文本被 word 分散到多个 run 的情况，重建段落并继承第一个 run 的样式
            for k in sorted_keys:
                if k in paragraph.text:
                    temp_txt = paragraph.text
                    for key in sorted_keys:
                        if key in temp_txt:
                            temp_txt = temp_txt.replace(key, str(replacements[key]))
                    
                    if len(paragraph.runs) > 0:
                        ref_run = paragraph.runs[0]
                        font_name = ref_run.font.name
                        font_size = ref_run.font.size
                        bold = ref_run.font.bold
                        italic = ref_run.font.italic
                        underline = ref_run.font.underline
                        color_rgb = ref_run.font.color.rgb if ref_run.font.color else None
                        
                        east_asia = None
                        if ref_run.element.rPr is not None and ref_run.element.rPr.rFonts is not None:
                            east_asia = ref_run.element.rPr.rFonts.get(qn('w:eastAsia'))
                        
                        paragraph.clear()
                        new_run = paragraph.add_run(temp_txt)
                        new_run.font.name = font_name
                        new_run.font.size = font_size
                        new_run.font.bold = bold
                        new_run.font.italic = italic
                        new_run.font.underline = underline
                        if color_rgb:
                            new_run.font.color.rgb = color_rgb
                        
                        if east_asia:
                            new_run.element.get_or_add_rPr().get_or_add_rFonts().set(qn('w:eastAsia'), east_asia)
                        elif font_name:
                            new_run.element.get_or_add_rPr().get_or_add_rFonts().set(qn('w:eastAsia'), font_name)
                    else:
                        paragraph.clear()
                        paragraph.add_run(temp_txt)
                    break

        for p in doc.paragraphs:
            replace_run_text_preserve_format(p)

        for t in doc.tables:
            for r in t.rows:
                for c in r.cells:
                    for p in c.paragraphs:
                        replace_run_text_preserve_format(p)

        fname = f"致嘉_成交确认书_{get_file_naming_date_str(invoice_data['date'])}.docx"
        fpath = os.path.join(output_dir, fname)
        doc.save(fpath)
        return fpath
    except Exception:
        traceback.print_exc()
        return None


# --- 生成出口报关单 (Excel) ---
def create_export_declaration(invoice_data, template_path, output_dir, user_inputs):
    if not os.path.exists(template_path):
        return None
    try:
        wb = load_workbook(template_path)
        ws = wb.active
        data = invoice_data["data"]

        consignee = user_inputs.get("consignee", "")
        trade_country = user_inputs.get("trade_country", "")
        pack_type = user_inputs.get("pack_type", "")

        freight = user_inputs.get("freight", "")
        insurance = user_inputs.get("insurance", "")
        other_fees = user_inputs.get("other_fees", "")

        gross_weight = invoice_data.get("gross_weight", 0)
        net_weight = invoice_data.get("net_weight", 0)
        
        total_packages = invoice_data.get("total_packages", len(data))

        incoterms = user_inputs.get("incoterms", "CIF")
        currency_type = user_inputs.get("currency", "EUR")

        curr_name = "美元" if currency_type == "USD" else "欧元"
        currency_symbol = "$" if currency_type == "USD" else "€"

        unified_font = Font(name="Microsoft YaHei", size=11, bold=False)
        thin_border = Border(
            left=Side(style="thin"),
            right=Side(style="thin"),
            top=Side(style="thin"),
            bottom=Side(style="thin"),
        )

        replacements = {
            "Gestamp Baires S.A.": consignee,
            "ZTD20260202001": invoice_data["contract_number"],
            "中国香港": trade_country,
            "胶合板箱": pack_type,
            "件数:6件": f"件数:{total_packages}件",
            "706.5": f"{gross_weight:.2f}",
            "580": f"{net_weight:.2f}",
            "毛重: 706.5 KGS": f"毛重: {gross_weight:.2f} KGS",
            "净重: 580 KGS": f"净重: {net_weight:.2f} KGS",
            "FOB": incoterms,
            "TFEE": freight,
            "IFEE": insurance,
            "OFEE": other_fees,
        }

        for row in ws.iter_rows():
            for cell in row:
                if cell.value is None:
                    continue

                cell_value_str = str(cell.value)
                original_value = cell_value_str
                new_value = original_value

                for k, v in replacements.items():
                    if k in new_value:
                        new_value = new_value.replace(k, str(v))

                if new_value != original_value:
                    cell.value = new_value
                    cell.font = unified_font

        start_row = 11
        max_item_rows_in_template = 30
        num_items = len(data)
        col_idx = {"no": 1, "code": 2, "name": 3, "qty": 4, "unit": 5, "price": 6, "amt": 7, "curr": 8, "dest": 10}

        for i in range(num_items):
            current_row = start_row + i
            item = data[i]

            c_name = str(item[0]) if item[0] else ""

            if "\n" in c_name:
                parts = [p.strip() for p in c_name.split("\n") if p.strip()]
                disp_name = parts[1] if len(parts) > 1 and re.search(r'[\u4e00-\u9fa5]', parts[1]) else parts[0]
            else:
                disp_name = c_name

            cat = get_item_category(c_name)
            hs_code = ITEM_DECLARATION_TEMPLATES[cat]["code"] if cat and cat in ITEM_DECLARATION_TEMPLATES else ""

            qty_val = item[4] if item[4] else 0
            if isinstance(qty_val, float) and qty_val.is_integer():
                qty_val = int(qty_val)
            unit_val = get_unit_str(c_name)

            price_val = item[7] if item[7] else 0.0
            amt_val = item[8] if item[8] else 0.0

            def write_cell_with_style(col, val, apply_border=False):
                cell = ws.cell(row=current_row, column=col)
                cell.value = val
                cell.font = unified_font
                cell.alignment = Alignment(horizontal="center", vertical="center")
                if apply_border:
                    cell.border = thin_border

            write_cell_with_style(col_idx["no"], i + 1)
            write_cell_with_style(col_idx["code"], hs_code)
            write_cell_with_style(col_idx["name"], disp_name, apply_border=True)
            write_cell_with_style(col_idx["qty"], qty_val)
            write_cell_with_style(col_idx["unit"], unit_val)
            write_cell_with_style(col_idx["price"], f"{price_val:.2f}")
            write_cell_with_style(col_idx["amt"], f"{currency_symbol}{amt_val:.2f}")
            write_cell_with_style(col_idx["curr"], curr_name)
            write_cell_with_style(col_idx["dest"], invoice_data["destination"])

        rows_to_delete = max_item_rows_in_template - num_items
        if rows_to_delete > 0:
            # ============== 彻底修复：隐藏而不是删除多余空行，完美保留底部格式 ==============
            for r_idx in range(start_row + num_items, start_row + max_item_rows_in_template):
                ws.row_dimensions[r_idx].hidden = True
                for c_idx in range(1, 15):
                    ws.cell(row=r_idx, column=c_idx).value = None

        fname = f"致嘉_出口报关单_{get_file_naming_date_str(invoice_data['date'])}.xlsx"
        fpath = os.path.join(output_dir, fname)
        wb.save(fpath)
        return fpath

    except Exception:
        traceback.print_exc()
        return None


def _build_zip(output_dir, zip_name):
    zip_base = os.path.join(tempfile.gettempdir(), zip_name.replace(".zip", ""))
    zip_path = shutil.make_archive(zip_base, "zip", output_dir)
    with open(zip_path, "rb") as f:
        data = f.read()
    os.remove(zip_path)
    return data


def render():
    st.header("东莞致嘉 · 出口单证自动生成")
    st.caption("上传 Excel 发票，自动生成：申报要素(Word)、成交确认书(Word)、出口报关单(Excel)，打包为一个 ZIP 下载。")

    uploaded = st.file_uploader("1. 上传 Excel 发票文件 (.xlsx / .xls)", type=["xlsx", "xls"], key="zhijia_upload")

    st.markdown("**2. 填写单证信息**")
    
    c1, c2 = st.columns(2)
    buyer_name_raw = c1.text_input("买方名称： (成交确认书)", key="zhijia_buyer", placeholder="默认: 香港致达五金制品有限公司")
    consignee = c2.text_input("境外收货人 (出口报关单)", key="zhijia_consignee", placeholder="例如: ABC Company")
    
    trade_country = c1.text_input("贸易国/目的国", key="zhijia_country", placeholder="例如: 德国")
    pack_type = c2.text_input("包装种类", value="胶合板箱", key="zhijia_pack")
    
    incoterms = c1.selectbox("成交方式", ["CIF", "FOB", "EXW"], key="zhijia_incoterms")
    currency = c2.selectbox("货币类型", ["EUR", "USD"], format_func=lambda x: "欧元 (EUR)" if x == "EUR" else "美元 (USD)", key="zhijia_currency")

    c3, c4, c5 = st.columns(3)
    freight = c3.text_input("运费 (可选)", key="zhijia_freight")
    insurance = c4.text_input("保费 (可选)", key="zhijia_insurance")
    other_fees = c5.text_input("杂费 (可选)", key="zhijia_fees")

    if st.button("生成文档", type="primary", key="zhijia_btn"):
        st.session_state.pop("zhijia_zip", None)
        st.session_state.pop("zhijia_summary", None)

        if uploaded is None:
            st.error("请先上传 Excel 发票文件。")
            return
        if not consignee.strip():
            st.error("请填写【境外收货人】。")
            return
        if not trade_country.strip():
            st.error("请填写贸易国/目的国。")
            return

        with st.spinner("正在解析发票并生成单证，请稍候…"):
            tmp = tempfile.mkdtemp()
            try:
                excel_path = os.path.join(tmp, uploaded.name)
                with open(excel_path, "wb") as f:
                    f.write(uploaded.getbuffer())

                invoice_data = read_invoice_data(excel_path)
                if not invoice_data or not invoice_data.get("data"):
                    st.error("无法读取发票数据或发票数据为空，请检查文件内容与列名（中文品名/型号/数量/单价/金额）。")
                    return

                num_items = len(invoice_data["data"])
                st.session_state["zhijia_summary"] = {
                    "公司": invoice_data["company_name"],
                    "日期": invoice_data["date"],
                    "目的国": invoice_data["destination"],
                    "合同号": invoice_data["contract_number"],
                    "商品数": num_items,
                    "总金额": invoice_data["total_amount"],
                    "总件数": invoice_data["total_packages"],
                    "净重": invoice_data["net_weight"],
                    "毛重": invoice_data["gross_weight"],
                    "found_weights": invoice_data["found_weights"],
                }
                
                buyer_name = buyer_name_raw.strip() if buyer_name_raw.strip() else "香港致达五金制品有限公司"

                user_inputs = {
                    "buyer_name": buyer_name,
                    "incoterms": incoterms,
                    "consignee": consignee,
                    "trade_country": trade_country,
                    "pack_type": pack_type,
                    "currency": currency,
                    "freight": freight,
                    "insurance": insurance,
                    "other_fees": other_fees,
                }

                path1 = create_declaration_elements(invoice_data, tmp)

                tpl_conf_name = "成交确认书模版1.docx"
                if num_items <= 8:
                    tpl_conf_name = "成交确认书模版1.docx"
                elif num_items <= 18:
                    tpl_conf_name = "成交确认书模版2.docx"
                elif num_items <= 30:
                    tpl_conf_name = "成交确认书模版3.docx"
                else:
                    tpl_conf_name = "成交确认书模版3.docx"
                tpl_conf = os.path.join(STATIC_DIR, tpl_conf_name)
                path2 = modify_sales_confirmation(invoice_data, tpl_conf, tmp, user_inputs) if os.path.exists(tpl_conf) else None

                tpl_decl = os.path.join(STATIC_DIR, "新版出口报关单样板.xlsx")
                path3 = create_export_declaration(invoice_data, tpl_decl, tmp, user_inputs) if os.path.exists(tpl_decl) else None

                files = [p for p in (path1, path2, path3) if p and os.path.exists(p)]
                if not files:
                    st.error("未能生成任何文件，请检查发票内容或模板文件。")
                    return

                zip_data = _build_zip(tmp, f"致嘉_单证_{get_file_naming_date_str(invoice_data['date'])}.zip")
                st.session_state["zhijia_zip"] = (f"单证_{get_file_naming_date_str(invoice_data['date'])}.zip", zip_data)
            finally:
                shutil.rmtree(tmp, ignore_errors=True)

    summary = st.session_state.get("zhijia_summary")
    if summary:
        st.success(
            f"✅ 读取成功：公司【{summary['公司']}】 合同号【{summary['合同号']}】\n\n"
            f"✅ 商品数： {summary['商品数']} 项\n\n"
            f"✅ 总金额： {summary['总金额']:,.2f} (通过'总价'和'总计'交点提取)\n\n"
            f"✅ 总件数： {summary['总件数']} (通过装箱单'箱数'和'总计'交点提取)"
        )
        if summary["found_weights"]:
            st.info(f"已从发票提取：净重 {summary['净重']:.2f} KG，毛重 {summary['毛重']:.2f} KG")
        else:
            st.warning("未在发票中找到“总计/Total”行，毛重/净重默认为 0。")

    zipdata = st.session_state.get("zhijia_zip")
    if zipdata:
        name, data = zipdata
        st.download_button("⬇ 下载文件包 (.zip)", data=data, file_name=name, mime="application/zip", key="zhijia_dl")


def main():
    render()


if __name__ == "__main__":
    main()
