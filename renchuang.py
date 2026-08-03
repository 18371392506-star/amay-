# coding: utf-8
"""宜章仁创 - 液晶显示屏单证生成（Streamlit 版）"""
import os
import re
import shutil
import tempfile
import traceback

import pandas as pd
import streamlit as st
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt
from openpyxl import load_workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.styles.colors import Color

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
STATIC_DIR = os.path.join(BASE_DIR, "static", "renchuang")


# --- Helper Functions ---
def get_file_naming_date_str():
    return _now_str()

def _now_str():
    from datetime import datetime
    return datetime.now().strftime("%Y%m%d")


def is_number(s):
    try:
        float(s)
        return True
    except (ValueError, TypeError):
        return False


def clean_decimal(val):
    try:
        if pd.isna(val) or str(val).strip() == "":
            return 0.0
        s = str(val).replace(",", "").replace("$", "").replace("￥", "").strip()
        return round(float(s), 2)
    except Exception:
        return 0.0


def read_renchuang_data(file_path):
    try:
        xls = pd.ExcelFile(file_path)
        sheet_names = xls.sheet_names

        packing_sheet_name = next((s for s in sheet_names if "装箱单" in s), None)
        if not packing_sheet_name:
            packing_sheet_name = next((s for s in sheet_names if "Packing" in s), None)
            if not packing_sheet_name:
                if sheet_names:
                    packing_sheet_name = sheet_names[0]
                else:
                    raise ValueError("No sheets found in Excel file.")

        df_pack = pd.read_excel(file_path, sheet_name=packing_sheet_name, header=None)

        contract_sheet_name = next((s for s in sheet_names if "购销合同" in s or "合同" in s), None)
        df_contract = None
        if contract_sheet_name:
            df_contract = pd.read_excel(file_path, sheet_name=contract_sheet_name, header=None)

        extracted_data = {
            "items": [],
            "contract_no": "",
            "destination": "中国香港",
            "total_packages": 0,
            "total_gross_weight": 0.0,
            "total_net_weight": 0.0,
        }

        if df_contract is not None:
            for r in range(len(df_contract)):
                for c in range(len(df_contract.columns)):
                    val = str(df_contract.iloc[r, c]).strip()
                    if "合同编号" in val:
                        parts = re.split(r"[:：]", val)
                        if len(parts) > 1 and parts[1].strip():
                            extracted_data["contract_no"] = parts[1].strip()
                        elif c + 1 < len(df_contract.columns):
                            extracted_data["contract_no"] = str(df_contract.iloc[r, c + 1]).strip()
                        break
                if extracted_data["contract_no"]:
                    break

            for r in range(len(df_contract)):
                for c in range(len(df_contract.columns)):
                    val = str(df_contract.iloc[r, c]).strip()
                    if "目的地" in val:
                        parts = re.split(r"[:：]", val)
                        if len(parts) > 1 and parts[1].strip():
                            extracted_data["destination"] = parts[1].strip()
                        elif c + 1 < len(df_contract.columns):
                            extracted_data["destination"] = str(df_contract.iloc[r, c + 1]).strip()
                        break
                if extracted_data["destination"] != "中国香港":
                    break

        header_row_idx = -1
        col_map = {}
        keywords = {
            "no": ["序号", "No."],
            "qty": ["数量", "QTY"],
            "unit": ["单位"],
            "nw": ["净重", "N.W"],
            "gw": ["毛重", "G.W"],
            "price": ["单价", "Unit Price", "单    价"],
            "total": ["总值", "Total", "金额", "Amount", "总   值"],
            "packages": ["包装件数", "件数", "CTNS"],
            "model": ["客户型号", "Model", "客户"],
            "size": ["尺寸", "Size"],
        }

        for r in range(min(20, len(df_pack))):
            row_vals_str_list = [str(val).strip() for val in df_pack.iloc[r] if pd.notna(val)]

            if any(k in val for val in row_vals_str_list for k in keywords["no"]):
                header_row_idx = r
                for c, val in enumerate(df_pack.iloc[r]):
                    val_str = str(val).strip() if pd.notna(val) else ""
                    for key, key_list in keywords.items():
                        if any(k_target == val_str or k_target in val_str for k_target in key_list):
                            col_map[key] = c
                            break
                break

        if header_row_idx == -1:
            raise ValueError("无法找到装箱单的表头，请确保包含'序号'等关键列。")

        start_row = header_row_idx + 1
        current_seq = 1

        def get_val(row_idx, col_key):
            if col_key in col_map and row_idx < len(df_pack) and col_map[col_key] < len(df_pack.columns):
                return df_pack.iloc[row_idx, col_map[col_key]]
            return None

        for r in range(start_row, len(df_pack)):
            no_val = get_val(r, "no")

            try:
                if pd.isna(no_val) or not is_number(no_val):
                    continue
                qty_raw = get_val(r, "qty")
                if not is_number(qty_raw):
                    continue
            except Exception:
                continue

            qty = int(clean_decimal(qty_raw))
            nw = clean_decimal(get_val(r, "nw"))
            gw = clean_decimal(get_val(r, "gw"))
            price = clean_decimal(get_val(r, "price"))
            total = clean_decimal(get_val(r, "total"))
            pkgs = int(clean_decimal(get_val(r, "packages")))

            model = str(get_val(r, "model")).strip() if get_val(r, "model") is not None else ""
            size = str(get_val(r, "size")).strip() if get_val(r, "size") is not None else ""

            if qty > 0:
                extracted_data["total_packages"] += pkgs
                extracted_data["total_gross_weight"] += gw
                extracted_data["total_net_weight"] += nw

                extracted_data["items"].append(
                    {
                        "seq": current_seq,
                        "qty": qty,
                        "unit": "个",
                        "nw": nw,
                        "gw": gw,
                        "price": price,
                        "total": total,
                        "packages": pkgs,
                        "model": model,
                        "size": size,
                    }
                )
                current_seq += 1

        return extracted_data

    except Exception as e:
        error_message = str(e)
        if "No engine for file type" in error_message or "xlrd is required" in error_message.lower():
            raise Exception("无法读取 .xls 文件。请确保已安装 'xlrd' 库或使用 .xlsx 格式文件。")
        elif "Unsupported format, or corrupt file: Expected BOF record" in error_message:
            raise Exception(f"Excel文件格式错误或文件已损坏: {error_message}")
        else:
            raise Exception(f"读取Excel数据失败: {error_message}")


# --- Generate Word Document (Declaration Elements) ---
def create_renchuang_declaration_elements(data, benefit_type, output_dir):
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(10.5)
    style.element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "Microsoft YaHei")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("申报要素")
    run.font.bold = True
    run.font.size = Pt(22)
    doc.add_paragraph()

    cn_nums = ["一", "二", "三", "四", "五", "六", "七", "八", "九", "十", "十一", "十二", "十三", "十四", "十五"]

    if not data["items"]:
        return None

    for i, item in enumerate(data["items"]):
        prefix = cn_nums[i] if i < len(cn_nums) else str(i + 1)

        p_head = doc.add_paragraph()
        run_head = p_head.add_run(f"{prefix}. 8524911000 液晶显示屏")
        run_head.font.bold = True
        run_head.font.size = Pt(12)

        size_str = f"{item['size']}英寸" if item["size"] else ""
        lines = [
            "1.品牌类型：",
            f"2. {benefit_type}",
            "3.用途：手机用",
            "4.液晶",
            "5.含驱动器及控制器",
            "6.不需要二次加工",
            f"7.尺寸:{size_str}",
            "8.无品牌",
            f"9.型号：{item['model']}",
            "10.非仅含GOA或类似栅极电路",
        ]

        p_body = doc.add_paragraph()
        p_body.add_run("\n".join(lines))

        doc.add_paragraph()

    fname = f"宜章仁创_申报要素_{get_file_naming_date_str()}.docx"
    fpath = os.path.join(output_dir, fname)
    doc.save(fpath)
    return fpath


# --- Generate Excel Document (Export Declaration) ---
def create_renchuang_export_declaration(data, template_path, output_dir):
    if not os.path.exists(template_path):
        return None

    try:
        wb = load_workbook(template_path)
        ws = wb.active

        item_data_start_row = 11
        min_data_col = 1
        max_data_col = 12

        template_footer_start_row = 12
        template_footer_end_row = 13
        num_footer_rows = template_footer_end_row - template_footer_start_row + 1

        template_footer_cells_data = []
        template_footer_row_heights = {}
        template_footer_merges_relative = []

        for r_idx in range(template_footer_start_row, template_footer_end_row + 1):
            template_footer_row_heights[r_idx] = ws.row_dimensions[r_idx].height
            row_cells_properties = []
            for c_idx in range(min_data_col, max_data_col + 1):
                cell = ws.cell(row=r_idx, column=c_idx)
                row_cells_properties.append(
                    {
                        "value": cell.value,
                        "font": cell.font.copy() if cell.font else None,
                        "border": cell.border.copy() if cell.border else None,
                        "alignment": cell.alignment.copy() if cell.alignment else None,
                        "fill": cell.fill.copy() if cell.fill else None,
                        "number_format": cell.number_format,
                    }
                )
            template_footer_cells_data.append(row_cells_properties)

        for merged_range in list(ws.merged_cells.ranges):
            if (
                merged_range.min_row >= template_footer_start_row
                and merged_range.max_row <= template_footer_end_row
                and merged_range.min_col >= min_data_col
                and merged_range.max_col <= max_data_col
            ):
                template_footer_merges_relative.append(
                    (
                        merged_range.min_col,
                        merged_range.max_col,
                        merged_range.min_row - template_footer_start_row,
                        merged_range.max_row - template_footer_start_row,
                    )
                )
                ws.unmerge_cells(str(merged_range))

        header_updates = [
            {"keyword": "件数", "cell_ref": "A1", "value_format": "件数:{total_packages}件", "requires_items": True},
            {"keyword": "毛重", "cell_ref": "A2", "value_format": "毛重（千克):{total_gross_weight:.2f}", "requires_items": True},
            {"keyword": "净重", "cell_ref": "A3", "value_format": "净重（千克）：{total_net_weight:.2f}", "requires_items": True},
            {"keyword": "合同协议号", "cell_ref": "A4", "value_format": "合同协议号:{contract_no}", "requires_items": False},
        ]

        for update_info in header_updates:
            keyword = update_info["keyword"]
            cell_ref = update_info["cell_ref"]
            value_format = update_info["value_format"]
            requires_items = update_info["requires_items"]

            current_data = {
                "total_packages": data["total_packages"],
                "total_gross_weight": data["total_gross_weight"],
                "total_net_weight": data["total_net_weight"],
                "contract_no": data["contract_no"],
            }

            if requires_items and len(data["items"]) == 0:
                if keyword == "件数":
                    formatted_value = "件数:0件"
                elif keyword == "毛重":
                    formatted_value = "毛重（千克):0.00"
                elif keyword == "净重":
                    formatted_value = "净重（千克）：0.00"
                else:
                    continue
            else:
                try:
                    formatted_value = value_format.format(**current_data)
                except Exception:
                    formatted_value = value_format.format(
                        total_packages=int(0), total_gross_weight=0.0, total_net_weight=0.0, contract_no=data["contract_no"]
                    )

            updated_directly = False
            try:
                cell_to_update = ws[cell_ref]
                if cell_to_update.value and isinstance(cell_to_update.value, str) and keyword in cell_to_update.value:
                    cell_to_update.value = formatted_value
                    updated_directly = True
            except Exception:
                pass

            if not updated_directly:
                found_in_search = False
                for r in range(1, 11):
                    for c_idx in range(1, ws.max_column + 1):
                        cell = ws.cell(row=r, column=c_idx)
                        if cell.value and isinstance(cell.value, str) and keyword in cell.value:
                            cell.value = formatted_value
                            found_in_search = True
                            break
                    if found_in_search:
                        break

        merged_ranges_to_unmerge = []
        for m_range in list(ws.merged_cells.ranges):
            if m_range.min_row >= item_data_start_row and m_range.max_col >= min_data_col and m_range.min_col <= max_data_col:
                merged_ranges_to_unmerge.append(m_range)
        for m_range_obj in merged_ranges_to_unmerge:
            ws.unmerge_cells(str(m_range_obj))

        ws.delete_rows(item_data_start_row + 1, amount=ws.max_row - item_data_start_row)

        num_items = len(data["items"])
        if num_items > 1:
            ws.insert_rows(item_data_start_row + 1, amount=num_items - 1)

        cols = {
            "seq": 1, "code": 2, "name": 3, "qty": 4, "unit": 5, "price": 6,
            "total": 7, "curr": 8, "origin": 9, "dest": 10, "source": 11, "tax": 12,
        }

        default_border = Border(
            left=Side(style="thin"), right=Side(style="thin"),
            top=Side(style="thin"), bottom=Side(style="thin"),
        )
        default_font = Font(name="Microsoft YaHei", size=10)
        align_center = Alignment(horizontal="center", vertical="center")

        default_row_height = ws.row_dimensions[item_data_start_row].height if ws.row_dimensions[item_data_start_row].height else 20

        for i, item in enumerate(data["items"]):
            current_row = item_data_start_row + i

            def write_cell_with_style(col_idx, val, alignment_obj):
                c = ws.cell(row=current_row, column=col_idx)
                c.value = val
                c.font = default_font
                c.alignment = alignment_obj
                c.border = default_border
                if col_idx in [cols["price"], cols["total"]]:
                    c.number_format = "#,##0.00"
                elif col_idx in [cols["qty"], cols["seq"]]:
                    c.number_format = "#,##0"
                ws.row_dimensions[current_row].height = default_row_height

            write_cell_with_style(cols["seq"], item["seq"], align_center)
            write_cell_with_style(cols["code"], "8524911000", align_center)
            write_cell_with_style(cols["name"], "液晶显示屏", align_center)
            write_cell_with_style(cols["qty"], item["qty"], align_center)
            write_cell_with_style(cols["unit"], "个", align_center)
            write_cell_with_style(cols["price"], item["price"], align_center)
            write_cell_with_style(cols["total"], item["total"], align_center)
            write_cell_with_style(cols["curr"], "USD", align_center)
            write_cell_with_style(cols["origin"], "中国", align_center)
            write_cell_with_style(cols["dest"], data["destination"], align_center)
            write_cell_with_style(cols["source"], "湖南郴州", align_center)
            write_cell_with_style(cols["tax"], "照章征免", align_center)

        new_footer_start_row = item_data_start_row + num_items

        ws.insert_rows(new_footer_start_row, amount=num_footer_rows)

        for i, row_cells_properties in enumerate(template_footer_cells_data):
            target_row = new_footer_start_row + i
            original_template_row = template_footer_start_row + i

            ws.row_dimensions[target_row].height = template_footer_row_heights[original_template_row]

            for c_idx_offset, cell_props in enumerate(row_cells_properties):
                target_col = min_data_col + c_idx_offset
                target_cell = ws.cell(row=target_row, column=target_col)
                target_cell.value = cell_props["value"]
                target_cell.font = cell_props["font"]
                target_cell.border = cell_props["border"]
                target_cell.alignment = cell_props["alignment"]
                target_cell.fill = cell_props["fill"]
                target_cell.number_format = cell_props["number_format"]

        for m_col_min, m_col_max, m_row_offset_start, m_row_offset_end in template_footer_merges_relative:
            new_min_row = new_footer_start_row + m_row_offset_start
            new_max_row = new_footer_start_row + m_row_offset_end
            ws.merge_cells(start_row=new_min_row, end_row=new_max_row, start_column=m_col_min, end_column=m_col_max)

        fname = f"宜章仁创_出口报关单_{get_file_naming_date_str()}.xlsx"
        fpath = os.path.join(output_dir, fname)
        wb.save(fpath)
        return fpath

    except Exception as e:
        traceback.print_exc()
        raise Exception(f"生成出口报关单失败: {e}")


def _build_zip(output_dir, zip_name):
    zip_base = os.path.join(tempfile.gettempdir(), zip_name.replace(".zip", ""))
    zip_path = shutil.make_archive(zip_base, "zip", output_dir)
    with open(zip_path, "rb") as f:
        data = f.read()
    os.remove(zip_path)
    return data


def render():
    st.header("宜章仁创 · 液晶显示屏出口单证生成")
    st.caption("上传 Excel（包含装箱单/购销合同），自动生成：申报要素(Word)、出口报关单(Excel)，打包为一个 ZIP 下载。")

    uploaded = st.file_uploader("1. 上传 Excel 文件 (.xlsx / .xls)", type=["xlsx", "xls"], key="rc_upload")

    benefit_type = st.selectbox("2. 享惠类型", ["不享惠", "享惠"], key="rc_benefit")

    if st.button("生成单证", type="primary", key="rc_btn"):
        st.session_state.pop("rc_zip", None)
        st.session_state.pop("rc_summary", None)

        if uploaded is None:
            st.error("请先上传 Excel 文件。")
            return

        with st.spinner("正在读取装箱单/合同并生成单证…"):
            tmp = tempfile.mkdtemp()
            try:
                excel_path = os.path.join(tmp, uploaded.name)
                with open(excel_path, "wb") as f:
                    f.write(uploaded.getbuffer())

                try:
                    data = read_renchuang_data(excel_path)
                except Exception as e:
                    st.error(str(e))
                    return

                if not data["items"]:
                    st.error("未在 Excel 中找到有效品名数据。请确保'装箱单'包含数字序号(1,2,3...)、数量、型号等信息。")
                    return

                st.session_state["rc_summary"] = {
                    "contract_no": data["contract_no"],
                    "destination": data["destination"],
                    "items": len(data["items"]),
                    "total_packages": data["total_packages"],
                    "total_gross_weight": data["total_gross_weight"],
                    "total_net_weight": data["total_net_weight"],
                }

                path_doc = create_renchuang_declaration_elements(data, benefit_type, tmp)
                tpl_path = os.path.join(STATIC_DIR, "新版出口报关单样板.xlsx")
                path_xls = create_renchuang_export_declaration(data, tpl_path, tmp) if os.path.exists(tpl_path) else None

                files = [p for p in (path_doc, path_xls) if p and os.path.exists(p)]
                if not files:
                    st.error("未能成功生成任何文件，请检查上传内容或模板。")
                    return

                zip_data = _build_zip(tmp, "宜章仁创_单证.zip")
                st.session_state["rc_zip"] = ("宜章仁创_单证.zip", zip_data)
            finally:
                shutil.rmtree(tmp, ignore_errors=True)

    summary = st.session_state.get("rc_summary")
    if summary:
        st.success(
            f"读取成功：{summary['items']} 个品名  合同号【{summary['contract_no'] or '未找到'}】 "
            f"目的国【{summary['destination']}】 "
            f"总件数 {summary['total_packages']} 件  毛重 {summary['total_gross_weight']:.2f} KG  净重 {summary['total_net_weight']:.2f} KG"
        )

    zipdata = st.session_state.get("rc_zip")
    if zipdata:
        name, data = zipdata
        st.download_button("⬇ 下载文件包 (.zip)", data=data, file_name=name, mime="application/zip", key="rc_dl")


def main():
    render()


if __name__ == "__main__":
    main()